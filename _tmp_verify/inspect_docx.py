"""Throwaway inspector: dump structure of a generated DOCX codebook."""
import sys
import zipfile
from docx import Document
from docx.oxml.ns import qn

path = sys.argv[1]
doc = Document(path)

section = doc.sections[0]
print('orientation:', section.orientation, 'page', section.page_width.mm, 'x', section.page_height.mm)
print('margins mm:', section.top_margin.mm, section.right_margin.mm, section.bottom_margin.mm, section.left_margin.mm)
print('title paragraph:', repr(doc.paragraphs[0].text))
print('normal font:', doc.styles['Normal'].font.name, doc.styles['Normal'].font.size.pt)

table = doc.tables[0]
print('columns:', len(table.columns), 'widths mm:', [round(c.width.mm, 1) for c in table.columns])
print('rows:', len(table.rows))

header = table.rows[0]
print('header cells:', [c.text for c in header.cells])
print('header repeats:', header._tr.find(qn('w:trPr')).find(qn('w:tblHeader')) is not None)

merged = 0
choice_tables = 0
for idx, row in enumerate(table.rows):
    tcs = row._tr.findall(qn('w:tc'))
    cells = row.cells
    if len(tcs) == 1:
        merged += 1
        print(f'  [row {idx}] SECTION BAND -> {cells[0].text!r}')
        continue
    nested = cells[2].tables
    if nested:
        choice_tables += 1
    field = cells[0].text
    question = cells[1].text.replace('\n', ' | ')
    answer = ' ; '.join(
        f"{r.cells[0].text}={r.cells[1].text}" for t in nested for r in t.rows
    )
    print(f'  [row {idx}] {field!r}')
    print(f'          Q: {question[:150]!r}')
    if answer:
        print(f'          A: {answer[:160]!r}')

print('merged section rows:', merged)
print('rows with choice tables:', choice_tables)

with zipfile.ZipFile(path) as zf:
    names = zf.namelist()
    xml = zf.read('word/document.xml').decode('utf-8')
print('package parts:', len(names))
print('document.xml bytes:', len(xml))
for token in ('w:tblHeader', 'w:cantSplit', 'w:shd', 'w:pBdr', 'w:tblLayout', 'w:szCs'):
    print(f'  {token} count:', xml.count('<' + token))
print('Arial Unicode MS refs:', xml.count('Arial Unicode MS'))
